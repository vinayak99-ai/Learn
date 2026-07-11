from docx import Document
import csv

def export_to_markdown(prd) -> str:
    lines = [f"# {prd.title}\n"]
    lines.append("## Problem Statement")
    lines.append(f"{prd.problem_statement}\n")

    lines.append("## Goals")
    for goal in prd.goals:
        lines.append(f"- {goal}")
    lines.append("")

    lines.append("## User Stories")
    for i, story in enumerate(prd.user_stories, 1):
        lines.append(f"### Story {i}")
        lines.append(f"**As a** {story.as_a}, **I want** {story.i_want}, "
                      f"**so that** {story.so_that}\n")
        lines.append("**Acceptance Criteria:**")
        for ac in story.acceptance_criteria:
            lines.append(f"- [ ] {ac}")
        lines.append("")

    lines.append("## Success Metrics")
    for metric in prd.success_metrics:
        lines.append(f"- {metric}")
    lines.append("")

    lines.append("## Assumptions")
    for a in prd.assumptions:
        lines.append(f"- {a}")

    return "\n".join(lines)


def export_to_docx(prd, path: str):
    doc = Document()
    doc.add_heading(prd.title, level=0)

    doc.add_heading("Problem Statement", level=1)
    doc.add_paragraph(prd.problem_statement)

    doc.add_heading("Goals", level=1)
    for goal in prd.goals:
        doc.add_paragraph(goal, style="List Bullet")

    doc.add_heading("User Stories", level=1)
    for i, story in enumerate(prd.user_stories, 1):
        doc.add_heading(f"Story {i}", level=2)
        p = doc.add_paragraph()
        p.add_run(f"As a {story.as_a}, I want {story.i_want}, "
                   f"so that {story.so_that}").italic = True

        doc.add_paragraph("Acceptance Criteria:", style="Intense Quote")
        for ac in story.acceptance_criteria:
            doc.add_paragraph(ac, style="List Bullet")

    doc.add_heading("Success Metrics", level=1)
    for m in prd.success_metrics:
        doc.add_paragraph(m, style="List Bullet")

    doc.add_heading("Assumptions", level=1)
    for a in prd.assumptions:
        doc.add_paragraph(a, style="List Bullet")

    doc.save(path)


def export_stories_to_csv(prd, path: str):
    """JIRA-importable CSV bridge (no API integration needed yet)."""
    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Summary", "Description", "Acceptance Criteria"])
        for story in prd.user_stories:
            summary = f"As a {story.as_a}, I want {story.i_want}"
            description = f"So that {story.so_that}"
            criteria = "; ".join(story.acceptance_criteria)
            writer.writerow([summary, description, criteria])
